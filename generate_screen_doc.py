# Developed by Saikiran Mekala
# email: mskvarma2020@gmail.com
# mobile: 9951456791
# Install python-docx module by running below command 
# pip install python-docx
import docx
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

class GenerateScreenDoc:
	"""
	GenerateScreenDoc Class
	==================================================================
	- This class helps you to add an image with a heading to the document.
	- Here I used python-docx module.
	- Before using this class please make sure you have installed python-docx module.
	"""
	doc = docx.Document()
	section = doc.sections[0]

	def __init__(self):
		"""
		Intializing default page settings.
		"""
		self.section.orientation = WD_ORIENT.LANDSCAPE
		self.section.page_height = Cm(21)
		self.section.page_width = Cm(29.7)
		self.section.top_margin = Cm(1.27)
		self.section.bottom_margin = Cm(1.27)
		self.section.left_margin = Cm(1.27)
		self.section.right_margin = Cm(1.27)

	
	def add_image_to_doc(self,image_url,heading=''):
		"""
		This method is useful for adding image to .docx file with a heading.
		Here heading is not manditory.
		@param1 : image_url 
		  type = string
		  example: './pic/01.png' i.e image name "01.png" in "pic" folder
		@param2 : heading
		  type = string
		  Note: This not a manditory parameter. Its upto you to pass this parameter. 
		  example: 'This is my heading'
		"""
		if len(heading)>0:
			self.doc.add_heading(heading, 3)
		assert len(image_url)>0,"Image Url should not be empty"
		self.doc.add_picture(image_url,width=Inches(10.7))

	def save_doc(self,file_name):
		"""
		This method used to save the document.
		@param1 : file_name
		  type = string
		  example: './reports/report1.docx' i.e stores report1.docx file in reports folder
		"""
		assert len(file_name)>0,"File name should not be empty" 
		self.doc.save(file_name)

